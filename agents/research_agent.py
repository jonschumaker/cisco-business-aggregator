#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
News Research Agent

This module provides functionality to research companies and generate comprehensive reports
based on news and web content. It integrates with Azure OpenAI, Tavily search, and Google Cloud Storage
to create well-structured reports in multiple formats (Markdown, Word, JSON).

Key Features:
- Deep research using AI models with web search capabilities
- Report generation in multiple formats (Markdown, Word, JSON)
- Integration with Google Cloud Storage for file storage
- Customer metadata integration from Excel files
- Human-readable formatting with proper styling and structure
"""

import os
from dotenv import load_dotenv
from typing import Union, List, Dict, Any, Optional
import uuid
import asyncio
import logging
from datetime import datetime
from urllib.parse import urlparse
import json
import pandas as pd
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import re
import docx
from google.cloud import storage
from google.oauth2 import service_account
import datetime as dt
import tempfile
import shutil
import openai
from tavily import TavilyClient

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables from .env file before any other imports
load_dotenv()

# Initialize API keys from environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

# Set up Google Cloud Storage credentials path
CREDENTIALS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "secrets", "google-credentials-dev.json")

# Set Azure OpenAI environment variables with the correct names
os.environ["OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY", OPENAI_API_KEY)
os.environ["OPENAI_API_VERSION"] = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview")
os.environ["AZURE_OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY")
os.environ["AZURE_OPENAI_ENDPOINT"] = os.getenv("AZURE_OPENAI_ENDPOINT", "https://phx-sales-ai.openai.azure.com/")
os.environ["AZURE_DEPLOYMENT_NAME"] = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")

# Check if API keys are available
if not os.environ.get("OPENAI_API_KEY"):
    logger.error("Neither OPENAI_API_KEY nor AZURE_OPENAI_API_KEY found in environment variables. Please add one of them to your .env file.")
    raise ValueError("Missing OpenAI API key in environment variables")
    
if not TAVILY_API_KEY:
    logger.error("TAVILY_API_KEY not found in environment variables. Please add it to your .env file.")
    raise ValueError("Missing TAVILY_API_KEY in environment variables")

# Set environment variables for libraries that need them
os.environ["TAVILY_API_KEY"] = TAVILY_API_KEY
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH

# Extract bucket information from env vars
GCS_BUCKET_PATH = os.getenv("OUTCOMES_PATH", "gs://sales-ai-dev-outcomes-6f1ce1c")
GCS_BUCKET_NAME = GCS_BUCKET_PATH.replace("gs://", "").split("/")[0]
GCS_FOLDER = "news-reports"
GCS_EXCEL_FOLDER = "data"

# Always use GCS for Excel files
USE_GCS_EXCEL = True

# Get local reports directory from environment variable or use default
LOCAL_REPORTS_DIR = os.getenv("LOCAL_REPORTS_DIR", "reports")
# Check if we should use local storage instead of GCS
USE_LOCAL_STORAGE = os.getenv("USE_LOCAL_STORAGE", "true").lower() in ["true", "1", "yes"]

# Module-level functions for backward compatibility
def get_temp_dir():
    """Create and return a temporary directory for storing files before uploading to GCS."""
    return tempfile.mkdtemp(prefix="news_reports_")

def get_gcs_client():
    """Get authenticated Google Cloud Storage client."""
    try:
        credentials = service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)
        return storage.Client(credentials=credentials)
    except Exception as e:
        logger.error(f"Error initializing GCS client: {str(e)}")
        return None

def extract_company_name(url: str, customer_name: Optional[str] = None) -> str:
    """
    Extract a clean company name from URL or customer name.
    
    Args:
        url: URL of the company website
        customer_name: Name of the customer
        
    Returns:
        str: Extracted company name
    """
    if customer_name:
        # Use customer name if provided
        return customer_name
    
    # Extract from URL
    try:
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        
        # Remove www. prefix if present
        if domain.startswith('www.'):
            domain = domain[4:]
        
        # Get the first part of the domain (before the first dot)
        company = domain.split('.')[0]
        
        # Clean up and capitalize
        company = company.replace('-', ' ').replace('_', ' ')
        company = ' '.join(word.capitalize() for word in company.split())
        
        return company
    except:
        # Fallback to a simple extraction
        clean_url = url.replace('http://', '').replace('https://', '').replace('www.', '')
        return clean_url.split('.')[0].capitalize()

async def process_url(url, topic=None, customer_name=None, customer_metadata=None, planner_model="gpt-4o", planner_provider="azure_openai"):
    """
    Process a single URL to generate a research report.
    
    Args:
        url: URL to research
        topic: Topic for the research
        customer_name: Name of the customer
        customer_metadata: Additional metadata about the customer
        planner_model: Model to use for planning
        planner_provider: Provider to use for planning
        
    Returns:
        Dict[str, str]: Dictionary with report file paths/URLs
    """
    # Create an instance of ResearchAgent
    agent = ResearchAgent()
    
    # Use the research_url method to generate a report
    return await agent.research_url(
        url=url,
        lookback_days=30,  # Default to 30 days
        customer_name=customer_name,
        customer_metadata=customer_metadata
    )

def load_customer_data():
    """
    Load customer data from Excel file and filter for valid websites and Heartland-Gulf.
    
    Returns:
        List[Tuple]: List of tuples with (website, customer_name, metadata_dict, topic)
    """
    try:
        # Define the Excel file path
        excel_filename = "Customer Parquet top 80 select hierarchy for test.xlsx"
        
        if USE_GCS_EXCEL:
            # If using GCS, download the Excel file from GCS
            logger.info("Using Excel file from Google Cloud Storage")
            
            # Create a temporary directory for the downloaded file
            temp_dir = get_temp_dir()
            local_excel_path = os.path.join(temp_dir, excel_filename)
            
            # GCS path to the Excel file
            gcs_excel_path = f"{GCS_EXCEL_FOLDER}/{excel_filename}"
            
            # Download the file
            download_success = download_from_gcs(gcs_excel_path, local_excel_path)
            
            if not download_success:
                logger.error(f"Failed to download Excel file from GCS. Falling back to local file.")
                # Fall back to the local file
                excel_path = excel_filename
            else:
                excel_path = local_excel_path
        else:
            # Use the local Excel file
            excel_path = excel_filename
        
        # Load the Excel file
        logger.info(f"Loading customer data from: {excel_path}")
        df = pd.read_excel(excel_path)
        
        logger.info(f"Loaded {len(df)} rows from Excel file")
        
        # Filter for valid websites (not blank or ".")
        df = df[df['WEBSITE'].notna()]  # Remove NaN values
        df = df[df['WEBSITE'] != "."]   # Remove "." values
        df = df[df['WEBSITE'] != ""]    # Remove empty strings
        
        # Filter for Heartland-Gulf in Sales Level 4
        df = df[df['SALES_LEVEL_4'] == 'HEARTLAND-GULF COMMERCIAL OPERATION']
        
        logger.info(f"Filtered to {len(df)} customers with valid websites and Heartland-Gulf region")
        
        # Create a list of tuples with (website, customer_name, metadata_dict, topic)
        customer_data = []
        for _, row in df.iterrows():
            website = row['WEBSITE'].strip() if isinstance(row['WEBSITE'], str) else row['WEBSITE']
            customer_name = row['SAVM_NAME_WITH_ID'] if 'SAVM_NAME_WITH_ID' in row else None
            
            # Create a dictionary with all available metadata from the Excel row
            metadata = {}
            for column in df.columns:
                # Convert any non-serializable values to strings
                if pd.notna(row[column]):
                    if isinstance(row[column], (int, float, str, bool)):
                        metadata[column] = row[column]
                    else:
                        metadata[column] = str(row[column])
            
            # Create a company-specific topic for each URL
            company_name = extract_company_name(website, customer_name)
            company_specific_topic = f"""Research and analyze news specifically about {company_name}.

The report structure should follow this exact order:

1. Company News (REQUIRES RESEARCH)
   - Any substantial news about the specific company {company_name}
   - Recent acquisitions, mergers, or partnerships involving {company_name}
   - Leadership changes or restructuring at {company_name}

2. News about their IT Priorities (REQUIRES RESEARCH)
   - {company_name}'s specific IT challenges, pain points, and desired outcomes
   - {company_name}'s technology investments or initiatives
   - {company_name}'s digital transformation efforts and IT strategy
   - Specific IT projects, implementations, or challenges faced by {company_name}

3. Discovery Questions for Cisco Sellers
   - Based on the findings in sections 1 and 2, create 3-5 thoughtful discovery questions that Cisco sellers can use to start conversations with {company_name}
   - Questions should relate to their IT challenges, pain points, or initiatives identified in the research
   - Questions should be open-ended and designed to uncover opportunities for Cisco solutions

Note: Focus ONLY on {company_name}. Do NOT include general industry trends or news about other companies unless they directly relate to {company_name}.
IMPORTANT: All sections REQUIRE research to find specific information about {company_name}.
"""
            
            customer_data.append((website, customer_name, metadata, company_specific_topic))
        
        return customer_data
    
    except Exception as e:
        logger.error(f"Error loading customer data: {str(e)}")
        return []

def download_from_gcs(gcs_path, local_path):
    """
    Download a file from Google Cloud Storage.
    
    Args:
        gcs_path: Path to the file in GCS
        local_path: Local path to save the file
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        client = get_gcs_client()
        if not client:
            logger.error("Failed to initialize GCS client")
            return False
            
        # Extract the blob path from the full GCS path
        if gcs_path.startswith("gs://"):
            # Remove the gs:// prefix and bucket name
            parts = gcs_path.replace("gs://", "").split("/", 1)
            if len(parts) < 2:
                logger.error(f"Invalid GCS path format: {gcs_path}")
                return False
            blob_path = parts[1]
        else:
            # Assume the path is already a blob path
            blob_path = gcs_path
        
        bucket = client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(blob_path)
        
        logger.info(f"Downloading file from gs://{GCS_BUCKET_NAME}/{blob_path} to {local_path}")
        blob.download_to_filename(local_path)
        
        if os.path.exists(local_path):
            logger.info(f"File successfully downloaded to {local_path}")
            return True
        else:
            logger.error(f"File download did not create the expected local file: {local_path}")
            return False
    except Exception as e:
        logger.error(f"Error downloading from GCS: {str(e)}")
        return False

class ResearchAgent:
    """
    Agent for researching companies and generating reports.
    
    This class provides methods for researching company URLs and generating
    comprehensive reports based on news and web content.
    """
    
    def __init__(self):
        """
        Initialize the research agent.
        """
        self.openai_client = self._initialize_openai_client()
        self.tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
    
    def _initialize_openai_client(self):
        """Initialize the OpenAI client with the appropriate configuration."""
        # Check if we're using Azure OpenAI
        if os.getenv("AZURE_OPENAI_API_KEY") and os.getenv("AZURE_OPENAI_ENDPOINT"):
            # Configure for Azure OpenAI
            return openai.AzureOpenAI(
                api_key=os.getenv("AZURE_OPENAI_API_KEY"),
                api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview"),
                azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
            )
        else:
            # Configure for regular OpenAI
            return openai.OpenAI(api_key=OPENAI_API_KEY)
    
    async def research_url(self, url: str, lookback_days: int = 30, 
                          customer_name: Optional[str] = None, 
                          customer_metadata: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        Research a URL and generate a report.
        
        Args:
            url: URL to research
            lookback_days: Number of days to look back for news
            customer_name: Name of the customer
            customer_metadata: Additional metadata about the customer
            
        Returns:
            Dict[str, str]: Dictionary with report file paths/URLs
        """
        logger.info(f"Researching URL: {url}")
        
        # Extract company name from URL or use customer name
        company_name = extract_company_name(url, customer_name)
        
        # Create a topic for research
        topic = f"Focus specifically on {company_name}, analyzing recent news about the company's operations, strategic initiatives, and IT priorities with emphasis on their specific IT challenges, pain points, and desired outcomes."
        
        try:
            # Search for information about the company
            search_results = await self._search_company_info(company_name, url, lookback_days)
            
            # If no results found, try with a longer lookback period
            if not search_results or len(search_results) < 3:
                logger.warning(f"Few or no results found with {lookback_days} days lookback. Trying with 90 days...")
                search_results = await self._search_company_info(company_name, url, 90)
                
                # If still no results, try with an even longer period
                if not search_results or len(search_results) < 3:
                    logger.warning(f"Few or no results found with 90 days lookback. Trying with 365 days...")
                    search_results = await self._search_company_info(company_name, url, 365)
            
            # Generate a report based on the search results
            if search_results and len(search_results) > 0:
                report_content = await self._generate_report(company_name, url, search_results, topic)
                
                # Save the report
                if report_content and len(report_content) > 200:
                    file_paths = self._save_markdown_report(
                        url=url,
                        content=report_content,
                        topic=topic,
                        customer_name=customer_name,
                        customer_metadata=customer_metadata
                    )
                    
                    logger.info(f"Research completed for {customer_name or url}")
                    return file_paths
                else:
                    logger.warning(f"No substantial report content was generated for {customer_name or url}")
                    return self._create_placeholder_report(url, company_name, customer_name, customer_metadata, lookback_days)
            else:
                logger.warning(f"No search results found for {customer_name or url}")
                return self._create_placeholder_report(url, company_name, customer_name, customer_metadata, lookback_days)
                
        except Exception as e:
            logger.error(f"Error researching {customer_name or url}: {str(e)}")
            return self._create_error_report(url, company_name, customer_name, customer_metadata, str(e))
    
    async def _search_company_info(self, company_name: str, url: str, lookback_days: int) -> List[Dict[str, Any]]:
        """
        Search for information about a company using Tavily.
        
        Args:
            company_name: Name of the company
            url: URL of the company website
            lookback_days: Number of days to look back for news
            
        Returns:
            List[Dict[str, Any]]: List of search results
        """
        try:
            # Create search queries
            queries = [
                f"{company_name} recent news",
                f"{company_name} IT priorities",
                f"{company_name} digital transformation",
                f"{company_name} technology investments",
                f"{company_name} IT challenges"
            ]
            
            all_results = []
            
            # Execute each search query
            for query in queries:
                try:
                    logger.info(f"Searching for: {query}")
                    
                    # Calculate the max_age_days parameter based on lookback_days
                    max_age_days = min(lookback_days, 365)  # Tavily has a maximum of 365 days
                    
                    # Execute the search
                    response = await asyncio.to_thread(
                        self.tavily_client.search,
                        query=query,
                        search_depth="advanced",
                        include_domains=[url] if url else None,
                        max_results=10,
                        max_age_days=max_age_days
                    )
                    
                    # Add results to the list
                    if response and "results" in response:
                        all_results.extend(response["results"])
                        logger.info(f"Found {len(response['results'])} results for query: {query}")
                    else:
                        logger.warning(f"No results found for query: {query}")
                        
                except Exception as search_error:
                    logger.error(f"Error searching for query '{query}': {str(search_error)}")
                    # Continue with other queries even if one fails
                    continue
            
            # Remove duplicates based on URL
            unique_results = []
            seen_urls = set()
            
            for result in all_results:
                if result["url"] not in seen_urls:
                    seen_urls.add(result["url"])
                    unique_results.append(result)
            
            logger.info(f"Found {len(unique_results)} unique results across all queries")
            return unique_results
            
        except Exception as e:
            logger.error(f"Error in search_company_info: {str(e)}")
            return []
    
    async def _generate_report(self, company_name: str, url: str, search_results: List[Dict[str, Any]], topic: str) -> str:
        """
        Generate a report based on search results using OpenAI.
        
        Args:
            company_name: Name of the company
            url: URL of the company website
            search_results: List of search results
            topic: Topic for the report
            
        Returns:
            str: Generated report content
        """
        try:
            # Prepare the context from search results
            context = self._prepare_context_from_results(search_results)
            
            # Create the prompt for the report
            prompt = f"""
            You are a business research analyst creating a report about {company_name}.
            
            Focus specifically on {company_name}, analyzing recent news about the company's operations, strategic initiatives, and IT priorities.
            The analysis should concentrate on {company_name}'s specific IT challenges, pain points, and desired outcomes,
            providing insights into their technology investments and digital transformation efforts.
            
            The report structure should follow this exact order:

            1. Company News (REQUIRES RESEARCH)
               - Any substantial news about the specific company {company_name}
               - Recent acquisitions, mergers, or partnerships involving {company_name}
               - Leadership changes or restructuring at {company_name}

            2. News about their IT Priorities (REQUIRES RESEARCH)
               - {company_name}'s specific IT challenges, pain points, and desired outcomes
               - {company_name}'s technology investments or initiatives
               - {company_name}'s digital transformation efforts and IT strategy
               - Specific IT projects, implementations, or challenges faced by {company_name}

            3. Discovery Questions for Cisco Sellers
               - Based on the findings in sections 1 and 2, create 3-5 thoughtful discovery questions that Cisco sellers can use to start conversations with {company_name}
               - Questions should relate to their IT challenges, pain points, or initiatives identified in the research
               - Questions should be open-ended and designed to uncover opportunities for Cisco solutions

            Note: Focus ONLY on {company_name}. Do NOT include general industry trends or news about other companies unless they directly relate to {company_name}.
            IMPORTANT: All sections REQUIRE research to find specific information about {company_name}.
            
            For each section, include relevant sources at the end of the section in the format:
            
            ### Sources
            [1]: URL1
            [2]: URL2
            
            Here is the research information:
            
            {context}
            """
            
            # Generate the report using OpenAI
            logger.info(f"Generating report for {company_name}")
            
            # Determine which model to use
            model = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")
            
            # Generate the report
            response = await asyncio.to_thread(
                self.openai_client.chat.completions.create,
                model=model,
                messages=[
                    {"role": "system", "content": "You are a business research analyst creating detailed reports about companies."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            # Extract the report content
            if response and response.choices and len(response.choices) > 0:
                report_content = response.choices[0].message.content
                logger.info(f"Generated report with {len(report_content)} characters")
                return report_content
            else:
                logger.warning("No content generated from OpenAI")
                return ""
                
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            return ""
    
    def _prepare_context_from_results(self, search_results: List[Dict[str, Any]]) -> str:
        """
        Prepare context from search results for the report generation.
        
        Args:
            search_results: List of search results
            
        Returns:
            str: Formatted context for the report
        """
        context = ""
        
        for i, result in enumerate(search_results):
            title = result.get("title", "No title")
            content = result.get("content", "No content")
            url = result.get("url", "No URL")
            
            context += f"Source {i+1}:\n"
            context += f"Title: {title}\n"
            context += f"URL: {url}\n"
            context += f"Content: {content}\n\n"
        
        return context
    
    def _standardize_sources_in_markdown(self, markdown_content: str) -> str:
        """
        Standardize source formatting in markdown content to use [number] format consistently.
        
        Args:
            markdown_content: Markdown content to standardize
            
        Returns:
            str: Standardized markdown content
        """
        lines = markdown_content.split('\n')
        result_lines = []
        in_sources_section = False
        source_counter = 0
        
        for line in lines:
            # Check if we're entering a sources section
            if line.strip() == "### Sources" or line.strip() == "Sources":
                in_sources_section = True
                source_counter = 0  # Reset counter for each sources section
                result_lines.append(line)
            # Check if we're leaving a sources section (new heading)
            elif in_sources_section and (line.startswith('#') or (len(result_lines) > 0 and result_lines[-1].strip() == "" and line.strip() == "")):
                in_sources_section = False
                result_lines.append(line)
            # Process source lines
            elif in_sources_section and line.strip():
                source_counter += 1
                
                # Check different source formats and standardize
                if line.strip().startswith('[') and ']:' in line:
                    # Format: [1]: URL - extract URL and reformat
                    source_parts = line.split(']: ', 1)
                    if len(source_parts) == 2:
                        url = source_parts[1]
                        result_lines.append(f"[{source_counter}]: {url}")
                    else:
                        result_lines.append(line)  # Keep as is if format is unexpected
                elif line.strip()[0].isdigit() and '. ' in line:
                    # Format: 1. URL - extract URL and reformat
                    source_parts = line.split('. ', 1)
                    if len(source_parts) == 2:
                        url = source_parts[1]
                        result_lines.append(f"[{source_counter}]: {url}")
                    else:
                        result_lines.append(line)  # Keep as is if format is unexpected
                else:
                    # Other format, just add bracketed number
                    result_lines.append(f"[{source_counter}]: {line.strip()}")
            else:
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def _save_json_report(self, markdown_content: str, output_path: str, url: str, topic: str, 
                         customer_name: Optional[str] = None, company_name: Optional[str] = None, 
                         customer_metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        Parse the markdown report into sections and save as a JSON file with metadata.
        
        Args:
            markdown_content: Markdown content to parse
            output_path: Path to save the JSON file
            url: URL of the company website
            topic: Topic of the report
            customer_name: Name of the customer
            company_name: Name of the company
            customer_metadata: Additional metadata about the customer
            
        Returns:
            str: Path to the saved JSON file
        """
        if not company_name and customer_name:
            company_name = extract_company_name(url, customer_name)
        elif not company_name:
            company_name = extract_company_name(url)
        
        # Common metadata for all sections
        common_metadata = {
            "url": url,
            "topic": topic,
            "customer_name": customer_name,
            "company_name": company_name,
            "generation_date": datetime.now().isoformat(),
        }
        
        # Primary key - SAVM_ID
        primary_key = None
        
        # Add Excel metadata if available
        if customer_metadata:
            # Include only the original column names (all caps)
            for key, value in customer_metadata.items():
                # Only add the original column name (which is typically all caps)
                common_metadata[key] = value
                
                # Set SAVM_ID as primary key if available
                if key == 'SAVM_ID':
                    primary_key = value
        
        # Parse the markdown content into sections
        sections = []
        
        # Use regex to find all headings and their content
        heading_pattern = re.compile(r'(#+)\s+(.*?)\n(.*?)(?=\n#+\s+|$)', re.DOTALL)
        matches = heading_pattern.findall(markdown_content)
        
        # Track the current parent section for attaching sources
        current_parent_section = None
        
        for i, (level, title, content) in enumerate(matches):
            # Check if this is a sources section
            is_sources_section = "sources" in title.lower()
            
            # If this is a sources section, attach it to the previous section and skip creating a new section
            if is_sources_section and current_parent_section is not None:
                # Extract URLs from the sources section
                urls = re.findall(r'https?://[^\s\]]+', content)
                current_parent_section["sources"] = urls
                continue
            
            # Determine section type based on title
            section_type = "other"
            if "introduction" in title.lower():
                section_type = "introduction"
            elif "company news" in title.lower():
                section_type = "company_news"
            elif "it priorities" in title.lower() or "digital transformation" in title.lower():
                section_type = "it_priorities"
            elif "discovery questions" in title.lower():
                section_type = "discovery_questions"
            
            # Create section object with minimal metadata - no duplication
            section = {
                "id": str(uuid.uuid4()),  # Unique identifier for each section
                "level": len(level),  # Number of # characters
                "title": title.strip(),
                "content": content.strip(),
                "section_type": section_type
            }
            
            # Add the section to our list
            sections.append(section)
            current_parent_section = section
        
        # Create the final JSON structure with SAVM_ID as primary key
        json_data = {
            "id": primary_key,  # Use SAVM_ID as the primary key
            "metadata": common_metadata,
            "sections": sections
        }
        
        # Save to JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"JSON report saved to file: {output_path}")
        return output_path
    
    def _markdown_to_word(self, markdown_content: str, output_path: str, title: str) -> str:
        """
        Convert markdown content to a well-formatted Word document.
        
        Args:
            markdown_content: Markdown content to convert
            output_path: Path to save the Word document
            title: Title for the Word document
            
        Returns:
            str: Path to the saved Word document
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a line break after title
            doc.add_paragraph()
            
            # Split the markdown content into lines
            lines = markdown_content.split('\n')
            
            # Process each line
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Skip empty lines
                if not line:
                    i += 1
                    continue
                
                # Process headings
                if line.startswith('#'):
                    # Count the number of # to determine heading level
                    level = 0
                    while level < len(line) and line[level] == '#':
                        level += 1
                    
                    # Extract the heading text
                    heading_text = line[level:].strip()
                    
                    # Add the heading with appropriate style
                    heading = doc.add_heading(heading_text, level=min(level, 9))
                    
                    i += 1
                    continue
                
                # Process lists
                if line.startswith('- ') or line.startswith('* '):
                    # Start a bulleted list
                    list_text = line[2:].strip()
                    p = doc.add_paragraph(list_text, style='List Bullet')
                    
                    i += 1
                    continue
                
                # Process numbered lists
                if re.match(r'^\d+\.\s', line):
                    # Start a numbered list
                    list_text = re.sub(r'^\d+\.\s', '', line).strip()
                    p = doc.add_paragraph(list_text, style='List Number')
                    
                    i += 1
                    continue
                
                # Process regular paragraphs
                paragraph = doc.add_paragraph()
                
                # Handle bold and italic formatting
                current_text = line
                
                # Replace bold markdown with Word bold
                bold_matches = re.finditer(r'\*\*(.*?)\*\*', current_text)
                last_end = 0
                new_text = ""
                
                for match in bold_matches:
                    new_text += current_text[last_end:match.start()]
                    last_end = match.end()
                    
                    # Add the bold text
                    run = paragraph.add_run(match.group(1))
                    run.bold = True
                
                # Add any remaining text
                if last_end < len(current_text):
                    paragraph.add_run(current_text[last_end:])
                
                i += 1
            
            # Save the document
            doc.save(output_path)
            logger.info(f"Word document saved to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            return ""
    
    def _save_markdown_report(self, url: str, content: str, topic: str, 
                             customer_name: Optional[str] = None, 
                             customer_metadata: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """
        Save a markdown report to file and convert to other formats.
        
        Args:
            url: URL of the company website
            content: Markdown content to save
            topic: Topic of the report
            customer_name: Name of the customer
            customer_metadata: Additional metadata about the customer
            
        Returns:
            Dict[str, str]: Dictionary with report file paths/URLs
        """
        try:
            # Extract company name from URL or use customer name
            company_name = extract_company_name(url, customer_name)
            
            # Standardize source formatting
            content = self._standardize_sources_in_markdown(content)
            
            # Create a unique filename based on company name and timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_company_name = re.sub(r'[^\w\-]', '_', company_name)
            base_filename = f"{safe_company_name}_{timestamp}"
            
            # Create the reports directory if it doesn't exist
            if USE_LOCAL_STORAGE:
                # Use local storage
                if not os.path.exists(LOCAL_REPORTS_DIR):
                    os.makedirs(LOCAL_REPORTS_DIR)
                
                # Define file paths
                md_path = os.path.join(LOCAL_REPORTS_DIR, f"{base_filename}.md")
                docx_path = os.path.join(LOCAL_REPORTS_DIR, f"{base_filename}.docx")
                json_path = os.path.join(LOCAL_REPORTS_DIR, f"{base_filename}.json")
                
                # Save markdown file
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Convert to Word document
                self._markdown_to_word(content, docx_path, f"Research Report: {company_name}")
                
                # Convert to JSON
                self._save_json_report(content, json_path, url, topic, customer_name, company_name, customer_metadata)
                
                # Return file paths
                return {
                    "markdown": md_path,
                    "word": docx_path,
                    "json": json_path
                }
            else:
                # Use Google Cloud Storage
                # Create a temporary directory for the files
                temp_dir = get_temp_dir()
                
                # Define temporary file paths
                temp_md_path = os.path.join(temp_dir, f"{base_filename}.md")
                temp_docx_path = os.path.join(temp_dir, f"{base_filename}.docx")
                temp_json_path = os.path.join(temp_dir, f"{base_filename}.json")
                
                # Save markdown file
                with open(temp_md_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Convert to Word document
                self._markdown_to_word(content, temp_docx_path, f"Research Report: {company_name}")
                
                # Convert to JSON
                self._save_json_report(content, temp_json_path, url, topic, customer_name, company_name, customer_metadata)
                
                # Upload files to GCS
                gcs_md_path = self._upload_to_gcs(temp_md_path, f"{GCS_FOLDER}/{base_filename}.md")
                gcs_docx_path = self._upload_to_gcs(temp_docx_path, f"{GCS_FOLDER}/{base_filename}.docx")
                gcs_json_path = self._upload_to_gcs(temp_json_path, f"{GCS_FOLDER}/{base_filename}.json")
                
                # Clean up temporary directory
                shutil.rmtree(temp_dir)
                
                # Return GCS URLs
                return {
                    "markdown": gcs_md_path,
                    "word": gcs_docx_path,
                    "json": gcs_json_path
                }
                
        except Exception as e:
            logger.error(f"Error saving markdown report: {str(e)}")
            return {}
    
    def _upload_to_gcs(self, local_path: str, gcs_blob_path: str) -> str:
        """
        Upload a file to Google Cloud Storage.
        
        Args:
            local_path: Local path of the file to upload
            gcs_blob_path: Path in GCS where the file should be stored
            
        Returns:
            str: GCS URL of the uploaded file
        """
        try:
            client = get_gcs_client()
            if not client:
                logger.error("Failed to initialize GCS client")
                return ""
                
            bucket = client.bucket(GCS_BUCKET_NAME)
            blob = bucket.blob(gcs_blob_path)
            
            logger.info(f"Uploading file from {local_path} to gs://{GCS_BUCKET_NAME}/{gcs_blob_path}")
            blob.upload_from_filename(local_path)
            
            # Return the GCS URL
            gcs_url = f"gs://{GCS_BUCKET_NAME}/{gcs_blob_path}"
            logger.info(f"File uploaded to {gcs_url}")
            return gcs_url
            
        except Exception as e:
            logger.error(f"Error uploading to GCS: {str(e)}")
            return ""
    
    def _create_placeholder_report(self, url: str, company_name: str, 
                                  customer_name: Optional[str] = None, 
                                  customer_metadata: Optional[Dict[str, Any]] = None,
                                  lookback_days: int = 30) -> Dict[str, str]:
        """
        Create a placeholder report when no search results are found.
        
        Args:
            url: URL of the company website
            company_name: Name of the company
            customer_name: Name of the customer
            customer_metadata: Additional metadata about the customer
            lookback_days: Number of days that were searched
            
        Returns:
            Dict[str, str]: Dictionary with report file paths/URLs
        """
        # Create a placeholder markdown content
        content = f"""# Research Report: {company_name}

## No Recent Information Found

After searching for information about {company_name} for the past {lookback_days} days, no substantial information was found. This could be due to:

- Limited online presence for {company_name}
- No recent news or updates published about the company
- The company may operate in a niche market with limited public information
- The company may be privately held with minimal public disclosures

## Recommendations

- Consider direct outreach to {company_name} to gather more information
- Explore industry-specific sources that might not be indexed by general search engines
- Connect with industry experts who might have insights about {company_name}
- Check specialized databases or subscription services for more detailed information

## Discovery Questions for Cisco Sellers

1. What are your current IT infrastructure challenges and priorities?
2. How is your organization approaching digital transformation initiatives?
3. What technology investments are you planning in the next 12-18 months?
4. How do you currently manage your network and security requirements?
5. What business outcomes are you looking to achieve with your technology investments?
"""
        
        # Create a topic for the report
        topic = f"Placeholder report for {company_name} due to insufficient information"
        
        # Save the report
        return self._save_markdown_report(
            url=url,
            content=content,
            topic=topic,
            customer_name=customer_name,
            customer_metadata=customer_metadata
        )
    
    def _create_error_report(self, url: str, company_name: str, 
                            customer_name: Optional[str] = None, 
                            customer_metadata: Optional[Dict[str, Any]] = None,
                            error_message: str = "") -> Dict[str, str]:
        """
        Create an error report when an exception occurs during research.
        
        Args:
            url: URL of the company website
            company_name: Name of the company
            customer_name: Name of the customer
            customer_metadata: Additional metadata about the customer
            error_message: Error message from the exception
            
        Returns:
            Dict[str, str]: Dictionary with report file paths/URLs
        """
        # Create an error markdown content
        content = f"""# Research Report: {company_name}

## Error Occurred During Research

An error occurred while researching information about {company_name}. The system was unable to complete the research process due to technical issues.

### Error Details

```
{error_message}
```

## Recommendations

- Try running the research again later
- Check if the company URL is correct: {url}
- Verify that all required API keys and credentials are properly configured
- If the issue persists, please contact technical support

## Alternative Research Methods

- Consider direct outreach to {company_name} to gather information
- Explore industry-specific sources manually
- Connect with industry experts who might have insights about {company_name}
"""
        
        # Create a topic for the report
        topic = f"Error report for {company_name} research"
        
        # Save the report
        return self._save_markdown_report(
            url=url,
            content=content,
            topic=topic,
            customer_name=customer_name,
            customer_metadata=customer_metadata
        )
