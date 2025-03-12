#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
News Research Agent

This module provides functionality to research companies and generate comprehensive reports
based on news and web content. It integrates with Azure OpenAI, Tavily search, and Google Cloud Storage
to create well-structured reports in multiple formats (Markdown, Word, JSON).

Key Features:
- Deep research using AI models with web search capabilities
- Report generation in multiple formats (Markdown, Word, JSON)
- Integration with Google Cloud Storage for file storage
- Customer metadata integration from Excel files
- Human-readable formatting with proper styling and structure
"""

import os
from dotenv import load_dotenv
from typing import Union, List
import uuid
import asyncio
import logging
from datetime import datetime
from urllib.parse import urlparse
import json
import pandas as pd
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time  # Add time module for sleep between retries
import re
import docx
from google.cloud import storage
from google.oauth2 import service_account
import datetime as dt  # Rename datetime module import to dt
import tempfile  # Add tempfile module for temporary file handling
import shutil  # For file operations

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables from .env file before any other imports
load_dotenv()

# Initialize API keys from environment variables
# No hardcoded keys - these should be set in your .env file
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

# Set up Google Cloud Storage credentials path
CREDENTIALS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "secrets", "google-credentials-dev.json")

# Set Azure OpenAI environment variables with the correct names
# This resolves the environment variable name mismatch issue
os.environ["OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY", OPENAI_API_KEY)  # Use Azure key for OpenAI
os.environ["OPENAI_API_VERSION"] = os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview")
os.environ["AZURE_OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY")
os.environ["AZURE_OPENAI_ENDPOINT"] = os.getenv("AZURE_OPENAI_ENDPOINT", "https://phx-sales-ai.openai.azure.com/")
os.environ["AZURE_DEPLOYMENT_NAME"] = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o")

# Check if API keys are available
if not os.environ.get("OPENAI_API_KEY"):
    logger.error("Neither OPENAI_API_KEY nor AZURE_OPENAI_API_KEY found in environment variables. Please add one of them to your .env file.")
    raise ValueError("Missing OpenAI API key in environment variables")
    
if not TAVILY_API_KEY:
    logger.error("TAVILY_API_KEY not found in environment variables. Please add it to your .env file.")
    raise ValueError("Missing TAVILY_API_KEY in environment variables")

# Set environment variables for libraries that need them
os.environ["TAVILY_API_KEY"] = TAVILY_API_KEY
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDENTIALS_PATH

# Extract bucket information from env vars
GCS_BUCKET_PATH = os.getenv("OUTCOMES_PATH", "gs://sales-ai-dev-outcomes-6f1ce1c")
GCS_BUCKET_NAME = GCS_BUCKET_PATH.replace("gs://", "").split("/")[0]
GCS_FOLDER = "news-reports"
GCS_EXCEL_FOLDER = "data"  # Folder where Excel files are stored in GCS

# Always use GCS for Excel files
USE_GCS_EXCEL = True

def get_temp_dir():
    """Create and return a temporary directory for storing files before uploading to GCS."""
    return tempfile.mkdtemp(prefix="news_reports_")

def get_gcs_client():
    """Get authenticated Google Cloud Storage client."""
    try:
        credentials = service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)
        return storage.Client(credentials=credentials)
    except Exception as e:
        logger.error(f"Error initializing GCS client: {str(e)}")
        return None

def upload_to_gcs(local_file_path, gcs_destination_path):
    """Upload a file to Google Cloud Storage."""
    try:
        client = get_gcs_client()
        if not client:
            logger.error("Failed to initialize GCS client")
            return None
            
        bucket = client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(gcs_destination_path)
        
        logger.info(f"Uploading file from {local_file_path} to GCS: gs://{GCS_BUCKET_NAME}/{gcs_destination_path}")
        blob.upload_from_filename(local_file_path)
        
        # Generate a signed URL that expires in 7 days
        url = blob.generate_signed_url(
            version="v4",
            expiration=dt.timedelta(days=7),
            method="GET"
        )
        
        logger.info(f"File successfully uploaded to GCS: gs://{GCS_BUCKET_NAME}/{gcs_destination_path}")
        logger.info(f"Generated signed URL (expires in 7 days): {url}")
        return url
    except Exception as e:
        logger.error(f"Error uploading to GCS: {str(e)}")
        return None

from langgraph.checkpoint.memory import MemorySaver
from open_deep_research.graph import builder
from langgraph.types import Command

# Configure the graph
memory = MemorySaver()
graph = builder.compile(checkpointer=memory)

# Remove duplicate reports directory creation
# REPORTS_DIR = "reports"
# os.makedirs(REPORTS_DIR, exist_ok=True)

async def research_topic(
    topic: str,
    urls: Union[str, List[str]] = None,
    tavily_days_back: int = 365, 
    max_search_depth: int = 1,
    writer_provider: str = "azure_openai",
    writer_model: str = "gpt-4o",
    planner_provider: str = "azure_openai",
    planner_model: str = "gpt-4o"
):
    """
    Research a topic using the deep research tool with OpenAI.
    Returns an async generator that yields events from the research process.
    """
    # Convert single URL to list
    if isinstance(urls, str):
        urls = [urls]
    
    thread = {"configurable": {
        "thread_id": str(uuid.uuid4()),
        "search_api": "tavily",
        "planner_provider": planner_provider,
        "planner_model": planner_model,
        "writer_provider": writer_provider,
        "writer_model": writer_model,
        "max_search_depth": max_search_depth,
        "tavily_days_back": tavily_days_back,
    }}
    
    # Add URLs to the thread if provided
    if urls:
        thread["configurable"]["urls"] = urls
    
    logger.info(f"Starting research on: {topic}")
    logger.info(f"Configuration: {thread['configurable']}")
    
    # Start the research with the topic and yield each event
    async for event in graph.astream({"topic": topic}, thread, stream_mode="updates"):
        logger.info(event)
        yield event
    
    # Approve the report plan and proceed to report generation
    async for event in graph.astream(Command(resume=True), thread, stream_mode="updates"):
        logger.info(event)
        yield event
    
def standardize_sources_in_markdown(markdown_content: str) -> str:
    """
    Standardize source formatting in markdown content to use [number] format consistently.
    This helps avoid formatting issues when converting to Word documents.
    """
    lines = markdown_content.split('\n')
    result_lines = []
    in_sources_section = False
    source_counter = 0
    
    for line in lines:
        # Check if we're entering a sources section
        if line.strip() == "### Sources" or line.strip() == "Sources":
            in_sources_section = True
            source_counter = 0  # Reset counter for each sources section
            result_lines.append(line)
        # Check if we're leaving a sources section (new heading)
        elif in_sources_section and (line.startswith('#') or (len(result_lines) > 0 and result_lines[-1].strip() == "" and line.strip() == "")):
            in_sources_section = False
            result_lines.append(line)
        # Process source lines
        elif in_sources_section and line.strip():
            source_counter += 1
            
            # Check different source formats and standardize
            if line.strip().startswith('[') and ']:' in line:
                # Format: [1]: URL - extract URL and reformat
                source_parts = line.split(']: ', 1)
                if len(source_parts) == 2:
                    url = source_parts[1]
                    result_lines.append(f"[{source_counter}]: {url}")
                else:
                    result_lines.append(line)  # Keep as is if format is unexpected
            elif line.strip()[0].isdigit() and '. ' in line:
                # Format: 1. URL - extract URL and reformat
                source_parts = line.split('. ', 1)
                if len(source_parts) == 2:
                    url = source_parts[1]
                    result_lines.append(f"[{source_counter}]: {url}")
                else:
                    result_lines.append(line)  # Keep as is if format is unexpected
            else:
                # Other format, just add bracketed number
                result_lines.append(f"[{source_counter}]: {line.strip()}")
        else:
            result_lines.append(line)
    
    return '\n'.join(result_lines)

def save_markdown_report(url: str, content: str, topic: str, customer_name: str = None, customer_metadata: dict = None):
    """Save the research report as a Markdown file and convert to Word document, then upload to GCS."""
    temp_dir = None
    try:
        # Create a temporary directory for storing files before uploading to GCS
        temp_dir = get_temp_dir()
        logger.info(f"Created temporary directory: {temp_dir}")
        
        # Create a filename based on the customer name or URL
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if customer_name:
            # Use customer name from SAVM_NAME_WITH_ID
            # Replace any characters that might be invalid in filenames
            safe_name = "".join([c if c.isalnum() or c in [' ', '_', '-'] else '_' for c in customer_name])
            base_filename = f"{safe_name}_{timestamp}"
        else:
            # Fallback to URL if customer name not provided
            parsed_url = urlparse(url)
            domain = parsed_url.netloc
            if domain.startswith('www.'):
                domain = domain[4:]  # Remove 'www.' prefix
            base_filename = f"OUTCOMES_{domain}_{timestamp}"
        
        # Markdown file path (temporary storage)
        md_file_path = os.path.join(temp_dir, f"{base_filename}.md")
        
        # Create markdown content with a more specific introduction
        company_name = extract_company_name(url, customer_name)
        markdown_content = f"# Research Report on {customer_name or url}\n\n"
        markdown_content += f"*Generated on {datetime.now().strftime('%B %d, %Y')}*\n\n"
        markdown_content += f"## URL: {url}\n\n"
        
        # Check if the content already has an Introduction section
        if not "# Introduction" in content and not "## Introduction" in content:
            markdown_content += f"# Introduction\n\n"
            markdown_content += f"This report focuses specifically on {company_name}, analyzing recent news about the company's operations, strategic initiatives, and IT priorities. "
            markdown_content += f"The analysis concentrates on {company_name}'s specific IT challenges, pain points, and desired outcomes, "
            markdown_content += f"providing insights into their technology investments and digital transformation efforts.\n\n"
        
        markdown_content += content
        
        # Standardize source formatting in the markdown content
        markdown_content = standardize_sources_in_markdown(markdown_content)
        
        # Save the markdown file to temporary directory
        with open(md_file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        logger.info(f"Markdown report saved to temporary file: {md_file_path}")
        
        # Save the JSON version with chunked sections to temporary directory
        json_file_path = os.path.join(temp_dir, f"{base_filename}.json")
        save_json_report(markdown_content, json_file_path, url, topic, customer_name, company_name, customer_metadata)
        logger.info(f"JSON report saved to temporary file: {json_file_path}")
        
        # Convert to Word document in temporary directory
        docx_file_path = os.path.join(temp_dir, f"{base_filename}.docx")
        try:
            markdown_to_word(markdown_content, docx_file_path, customer_name or url)
            logger.info(f"Word document saved to temporary file: {docx_file_path}")
        except Exception as e:
            logger.error(f"Error converting to Word: {str(e)}")
            docx_file_path = None
        
        # Upload files to Google Cloud Storage
        result = {}
        upload_success = False
        
        # Upload markdown file to GCS
        if os.path.exists(md_file_path):
            gcs_md_path = f"{GCS_FOLDER}/{base_filename}.md"
            md_url = upload_to_gcs(md_file_path, gcs_md_path)
            if md_url:
                result["markdown_gcs_url"] = md_url
                upload_success = True
            else:
                logger.error(f"Failed to upload markdown file to GCS: {md_file_path}")
        else:
            logger.error(f"Markdown file does not exist: {md_file_path}")
            
        # Upload JSON file to GCS
        if os.path.exists(json_file_path):
            gcs_json_path = f"{GCS_FOLDER}/{base_filename}.json"
            json_url = upload_to_gcs(json_file_path, gcs_json_path)
            if json_url:
                result["json_gcs_url"] = json_url
                upload_success = True
            else:
                logger.error(f"Failed to upload JSON file to GCS: {json_file_path}")
        else:
            logger.error(f"JSON file does not exist: {json_file_path}")
            
        # Upload Word document to GCS if available
        if docx_file_path and os.path.exists(docx_file_path):
            gcs_docx_path = f"{GCS_FOLDER}/{base_filename}.docx"
            docx_url = upload_to_gcs(docx_file_path, gcs_docx_path)
            if docx_url:
                result["docx_gcs_url"] = docx_url
                upload_success = True
            else:
                logger.error(f"Failed to upload Word document to GCS: {docx_file_path}")
        elif docx_file_path:
            logger.error(f"Word document does not exist: {docx_file_path}")
        
        if not upload_success:
            logger.error("No files were successfully uploaded to GCS")
            return None
            
        return result
    except Exception as e:
        logger.error(f"Error saving and uploading report: {str(e)}")
        return None
    finally:
        # Clean up temporary directory
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                logger.error(f"Error cleaning up temporary directory: {str(e)}")

def save_json_report(markdown_content: str, output_path: str, url: str, topic: str, 
                    customer_name: str = None, company_name: str = None, customer_metadata: dict = None):
    """
    Parse the markdown report into sections and save as a JSON file with metadata.
    Each section is chunked and tagged with appropriate metadata.
    Sources are attached as metadata to their parent sections rather than as separate sections.
    """
    if not company_name and customer_name:
        company_name = extract_company_name(url, customer_name)
    elif not company_name:
        company_name = extract_company_name(url)
    
    # Common metadata for all sections
    common_metadata = {
        "url": url,
        "topic": topic,
        "customer_name": customer_name,
        "company_name": company_name,
        "generation_date": datetime.now().isoformat(),
    }
    
    # Primary key - SAVM_ID
    primary_key = None
    
    # Add Excel metadata if available
    if customer_metadata:
        # Include only the original column names (all caps)
        for key, value in customer_metadata.items():
            # Only add the original column name (which is typically all caps)
            common_metadata[key] = value
            
            # Set SAVM_ID as primary key if available
            if key == 'SAVM_ID':
                primary_key = value
    
    # Parse the markdown content into sections
    sections = []
    
    # Use regex to find all headings and their content
    heading_pattern = re.compile(r'(#+)\s+(.*?)\n(.*?)(?=\n#+\s+|$)', re.DOTALL)
    matches = heading_pattern.findall(markdown_content)
    
    # Track the current parent section for attaching sources
    current_parent_section = None
    
    for i, (level, title, content) in enumerate(matches):
        # Check if this is a sources section
        is_sources_section = "sources" in title.lower()
        
        # If this is a sources section, attach it to the previous section and skip creating a new section
        if is_sources_section and current_parent_section is not None:
            # Extract URLs from the sources section
            urls = re.findall(r'https?://[^\s\]]+', content)
            current_parent_section["sources"] = urls
            continue
        
        # Determine section type based on title
        section_type = "other"
        if "introduction" in title.lower():
            section_type = "introduction"
        elif "company news" in title.lower():
            section_type = "company_news"
        elif "it priorities" in title.lower() or "digital transformation" in title.lower():
            section_type = "it_priorities"
        elif "discovery questions" in title.lower():
            section_type = "discovery_questions"
        
        # Create section object with minimal metadata - no duplication
        section = {
            "id": str(uuid.uuid4()),  # Unique identifier for each section
            "level": len(level),  # Number of # characters
            "title": title.strip(),
            "content": content.strip(),
            "section_type": section_type
        }
        
        # Add the section to our list
        sections.append(section)
        current_parent_section = section
    
    # Create the final JSON structure with SAVM_ID as primary key
    json_data = {
        "id": primary_key,  # Use SAVM_ID as the primary key
        "metadata": common_metadata,
        "sections": sections
    }
    
    # Save to JSON file
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Simple JSON report saved to temporary file: {output_path}")
    return output_path

def markdown_to_word(markdown_content: str, output_path: str, title: str):
    """Convert markdown content to a well-formatted Word document."""
    # Create a new Document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = f"Research Report on {title}"
    doc.core_properties.author = "Research Agent"
    
    # Split the markdown content into sections for better processing
    sections = []
    current_section = {"heading": None, "content": [], "sources": []}
    lines = markdown_content.split('\n')
    in_sources_section = False
    
    # First pass: organize content into sections
    for line in lines:
        if line.startswith('# ') or line.startswith('## '):
            # New main section
            if current_section["heading"]:
                sections.append(current_section)
            current_section = {"heading": line, "content": [], "sources": []}
            in_sources_section = False
        elif line.startswith('### Sources') or line.startswith('Sources'):
            # Start of sources section
            in_sources_section = True
            # Don't add this line to content
        elif in_sources_section and line.strip():
            # This is a source line, add to sources
            current_section["sources"].append(line.strip())
        elif line.strip() and current_section["heading"]:
            # Regular content
            current_section["content"].append(line)
    
    # Add the last section
    if current_section["heading"]:
        sections.append(current_section)
    
    # Second pass: write to Word document
    for section in sections:
        # Add section heading
        if section["heading"].startswith('# '):
            p = doc.add_paragraph()
            heading_text = section["heading"].replace('# ', '')
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(16)
            p.style = 'Heading 1'
            
            # Center the title if it's the main report title
            if heading_text.lower().startswith("research report") or title.lower() in heading_text.lower():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add date under the title
                date_paragraph = doc.add_paragraph()
                date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
                date_run.italic = True
                date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add a separator
                doc.add_paragraph()
        elif section["heading"].startswith('## '):
            p = doc.add_paragraph()
            heading_text = section["heading"].replace('## ', '')
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(14)
            p.style = 'Heading 2'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Special handling for Discovery Questions section
        is_discovery_section = section["heading"].replace('## ', '').strip().lower().startswith("discovery questions")
        
        # For Discovery Questions section, first collect all numbered items
        if is_discovery_section:
            # First, add the introductory paragraph(s)
            intro_text = []
            i = 0
            while i < len(section["content"]) and not re.match(r'^\d+\.', section["content"][i].strip()):
                if section["content"][i].strip():
                    intro_text.append(section["content"][i])
                i += 1
            
            if intro_text:
                intro_para = doc.add_paragraph()
                for line in intro_text:
                    intro_para.add_run(line)
            
            # Now collect and process all numbered items
            numbered_items = []
            current_item = None
            
            while i < len(section["content"]):
                line = section["content"][i].strip()
                
                # Check if this is a new numbered item
                match = re.match(r'^(\d+)\.(.+)', line)
                if match:
                    if current_item:
                        numbered_items.append(current_item)
                    number, text = match.groups()
                    current_item = {"number": int(number), "text": text.strip()}
                # Check if this is a continuation of the current item (indented text)
                elif line and current_item:
                    current_item["text"] += " " + line
                # Empty line - end of current item
                elif not line and current_item:
                    numbered_items.append(current_item)
                    current_item = None
                
                i += 1
            
            # Add the last item if there is one
            if current_item:
                numbered_items.append(current_item)
            
            # Now add all numbered items with consistent formatting
            for item in numbered_items:
                p = doc.add_paragraph(item["text"], style='List Number')
        else:
            # Process content with special handling for tables and numbered lists
            i = 0
            in_table = False
            table_rows = []
            
            while i < len(section["content"]):
                line = section["content"][i].strip()
                
                # Check for table header or separator
                if line.startswith('|') and line.endswith('|'):
                    # This might be a table row
                    if not in_table:
                        # Start of a new table
                        in_table = True
                        table_rows = [line]
                    else:
                        # Continue existing table
                        table_rows.append(line)
                    i += 1
                    continue
                elif in_table and not line:
                    # End of table (empty line after table)
                    # Process the collected table rows
                    if len(table_rows) >= 3:  # Need at least header, separator, and one data row
                        # Create Word table
                        # First, parse the table structure
                        header_row = table_rows[0]
                        columns = [col.strip() for col in header_row.split('|')[1:-1]]  # Remove first and last empty elements
                        num_columns = len(columns)
                        
                        # Create table with appropriate number of rows and columns
                        table = doc.add_table(rows=len(table_rows) - 1, cols=num_columns)  # -1 to exclude separator row
                        table.style = 'Table Grid'
                        
                        # Process each row (skip the separator row)
                        row_idx = 0
                        for table_row in [table_rows[0]] + table_rows[2:]:  # Header + data rows (skip separator)
                            cells = [cell.strip() for cell in table_row.split('|')[1:-1]]
                            for col_idx, cell_content in enumerate(cells):
                                if col_idx < num_columns:  # Ensure we don't exceed the number of columns
                                    table.cell(row_idx, col_idx).text = cell_content
                            row_idx += 1
                    
                    # Reset table state
                    in_table = False
                    table_rows = []
                    i += 1
                    continue
                elif in_table:
                    # We're in a table but this line doesn't match table format
                    # End the table and process this line normally
                    if len(table_rows) >= 3:
                        # Create Word table as above
                        header_row = table_rows[0]
                        columns = [col.strip() for col in header_row.split('|')[1:-1]]
                        num_columns = len(columns)
                        
                        table = doc.add_table(rows=len(table_rows) - 1, cols=num_columns)
                        table.style = 'Table Grid'
                        
                        row_idx = 0
                        for table_row in [table_rows[0]] + table_rows[2:]:
                            cells = [cell.strip() for cell in table_row.split('|')[1:-1]]
                            for col_idx, cell_content in enumerate(cells):
                                if col_idx < num_columns:
                                    table.cell(row_idx, col_idx).text = cell_content
                            row_idx += 1
                    
                    # Reset table state
                    in_table = False
                    table_rows = []
                    # Don't increment i, process this line normally
                elif line == "":
                    # Empty line
                    doc.add_paragraph()
                    i += 1
                elif line.startswith('### '):
                    # Subheading
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('### ', ''))
                    run.bold = True
                    run.font.size = Pt(12)
                    p.style = 'Heading 3'
                    i += 1
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph(line.replace('- ', '').replace('* ', ''), style='List Bullet')
                    i += 1
                elif re.match(r'^\d+\.\s', line):
                    # Regular numbered list item
                    match = re.match(r'(\d+)\.?\s*(.*)', line.strip())
                    if match:
                        number, text = match.groups()
                        p = doc.add_paragraph(text.strip(), style='List Number')
                    else:
                        doc.add_paragraph(line)
                    i += 1
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    
                    # Handle markdown formatting
                    if '**' in line:
                        parts = line.split('**')
                        for j, part in enumerate(parts):
                            if j % 2 == 0:  # Regular text
                                p.add_run(part)
                            else:  # Bold text
                                p.add_run(part).bold = True
                    elif '*' in line and not line.startswith('*'):
                        parts = line.split('*')
                        for j, part in enumerate(parts):
                            if j % 2 == 0:  # Regular text
                                p.add_run(part)
                            else:  # Italic text
                                p.add_run(part).italic = True
                    else:
                        p.add_run(line)
                    
                    i += 1
            
            # Check if we ended with a table
            if in_table and len(table_rows) >= 3:
                # Create Word table
                header_row = table_rows[0]
                columns = [col.strip() for col in header_row.split('|')[1:-1]]
                num_columns = len(columns)
                
                table = doc.add_table(rows=len(table_rows) - 1, cols=num_columns)
                table.style = 'Table Grid'
                
                row_idx = 0
                for table_row in [table_rows[0]] + table_rows[2:]:
                    cells = [cell.strip() for cell in table_row.split('|')[1:-1]]
                    for col_idx, cell_content in enumerate(cells):
                        if col_idx < num_columns:
                            table.cell(row_idx, col_idx).text = cell_content
                    row_idx += 1
        
        # Add sources if present
        if section["sources"]:
            # Add Sources heading
            p = doc.add_paragraph()
            run = p.add_run("Sources")
            run.bold = True
            run.font.size = Pt(12)
            p.style = 'Heading 3'
            
            # Collect and process sources
            for i, source in enumerate(section["sources"]):
                # Create a new paragraph for each source
                p = doc.add_paragraph()
                
                # Extract the reference number and URL
                if source.strip().startswith('[') and ']:' in source:
                    # Format: [1]: URL or [1]: [1] Title URL
                    source_parts = source.split(']: ', 1)
                    if len(source_parts) == 2:
                        ref_num = source_parts[0]
                        content = source_parts[1].strip()
                        
                        # Add the reference number
                        p.add_run(f"{ref_num}]: ")
                        
                        # Check if content starts with another reference number like [1]
                        content_match = re.match(r'^\[(\d+)\](.*?)https?://', content)
                        if content_match:
                            # Skip the duplicate reference number in the content
                            title_text = content_match.group(2).strip()
                            url_start = content.find('http')
                            if url_start > 0:
                                url = content[url_start:].strip()
                                
                                # Add the title text (without the duplicate reference number)
                                if title_text:
                                    p.add_run(title_text + " ")
                                
                                # Add the URL as a hyperlink
                                add_hyperlink(p, url, url, '0000FF', True)
                        else:
                            # No duplicate reference number, process normally
                            url_match = re.search(r'https?://\S+', content)
                            if url_match:
                                url = url_match.group(0)
                                before_url = content[:url_match.start()].strip()
                                
                                # Add any text before the URL
                                if before_url:
                                    p.add_run(before_url + " ")
                                
                                # Add the URL as a hyperlink
                                add_hyperlink(p, url, url, '0000FF', True)
                                
                                # Add any text after the URL
                                after_url = content[url_match.end():].strip()
                                if after_url:
                                    p.add_run(" " + after_url)
                            else:
                                # No URL found, just add the content
                                p.add_run(content)
                elif source.strip()[0].isdigit() and '. ' in source:
                    # Format: 1. URL or 1. [1] Title URL
                    source_parts = source.split('. ', 1)
                    if len(source_parts) == 2:
                        ref_num = source_parts[0]
                        content = source_parts[1].strip()
                        
                        # Add the reference number
                        p.add_run(f"[{ref_num}]: ")
                        
                        # Check if content starts with another reference number like [1]
                        content_match = re.match(r'^\[(\d+)\](.*?)https?://', content)
                        if content_match:
                            # Skip the duplicate reference number in the content
                            title_text = content_match.group(2).strip()
                            url_start = content.find('http')
                            if url_start > 0:
                                url = content[url_start:].strip()
                                
                                # Add the title text (without the duplicate reference number)
                                if title_text:
                                    p.add_run(title_text + " ")
                                
                                # Add the URL as a hyperlink
                                add_hyperlink(p, url, url, '0000FF', True)
                        else:
                            # No duplicate reference number, process normally
                            url_match = re.search(r'https?://\S+', content)
                            if url_match:
                                url = url_match.group(0)
                                before_url = content[:url_match.start()].strip()
                                
                                # Add any text before the URL
                                if before_url:
                                    p.add_run(before_url + " ")
                                
                                # Add the URL as a hyperlink
                                add_hyperlink(p, url, url, '0000FF', True)
                                
                                # Add any text after the URL
                                after_url = content[url_match.end():].strip()
                                if after_url:
                                    p.add_run(" " + after_url)
                            else:
                                # No URL found, just add the content
                                p.add_run(content)
                else:
                    # Other format, just add as is
                    p.add_run(f"[{i+1}]: ")
                    
                    # Try to extract a URL if present
                    url_match = re.search(r'https?://\S+', source)
                    if url_match:
                        url = url_match.group(0)
                        before_url = source[:url_match.start()]
                        if before_url:
                            p.add_run(before_url)
                        add_hyperlink(p, url, url, '0000FF', True)
                        after_url = source[url_match.end():].strip()
                        if after_url:
                            p.add_run(" " + after_url)
                    else:
                        # No URL found, just add the text
                        p.add_run(source)
            
            # Add space after sources
            doc.add_paragraph()
    
    # Save the document
    doc.save(output_path)
    return output_path

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """
    Add a hyperlink to a paragraph.
    
    :param paragraph: The paragraph to add the hyperlink to
    :param text: The text to display
    :param url: The URL to link to
    :param color: The color of the link (in hex, default is blue)
    :param underline: Whether to underline the link (default is True)
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Create a new run object (a wrapper over a run element)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    
    # Create a new run properties object
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add color if provided
    if color:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    
    # Add underline if specified
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

# Helper function to debug events
def debug_event(event, prefix=""):
    """Debug helper to print event structure"""
    logger.debug(f"{prefix}Event type: {type(event)}")
    if isinstance(event, dict):
        logger.debug(f"{prefix}Event keys: {list(event.keys())}")
        for key in event.keys():
            if isinstance(event[key], dict):
                logger.debug(f"{prefix}  {key} keys: {list(event[key].keys())}")
    logger.debug(f"{prefix}Event content: {event}")
    logger.debug(f"{prefix}---")

def extract_company_name(url: str, customer_name: str = None):
    """Extract a clean company name from URL or customer name."""
    if customer_name:
        # Use customer name if provided
        return customer_name
    
    # Extract from URL
    try:
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        
        # Remove www. prefix if present
        if domain.startswith('www.'):
            domain = domain[4:]
        
        # Get the first part of the domain (before the first dot)
        company = domain.split('.')[0]
        
        # Clean up and capitalize
        company = company.replace('-', ' ').replace('_', ' ')
        company = ' '.join(word.capitalize() for word in company.split())
        
        return company
    except:
        # Fallback to a simple extraction
        clean_url = url.replace('http://', '').replace('https://', '').replace('www.', '')
        return clean_url.split('.')[0].capitalize()

async def process_url(url, topic=None, customer_name=None, customer_metadata=None, planner_model="gpt-4o", planner_provider="azure_openai"):
    """Process a single URL to generate a research report"""
    logger.info(f"🔍 Researching {customer_name or url}...")
    
    # Extract company name from URL or use customer name
    company_name = extract_company_name(url, customer_name)
    
    # If no topic is provided, create a default one
    if not topic:
        topic = f"Focus specifically on {company_name}, analyzing recent news about the company's operations, strategic initiatives, and IT priorities with emphasis on their specific IT challenges, pain points, and desired outcomes."
    
    # Collect the final report content
    report_content = ""
    
    # First try with 60 days lookback
    tavily_days_back = 60
    max_retries = 3  # Maximum number of retries for API failures
    
    try:
        # Run the research with retry mechanism
        retry_count = 0
        while retry_count < max_retries:
            try:
                # Run the research
                async for event in research_topic(
                            topic=topic,  # Use the provided company-specific topic
                            urls=[url],
                            tavily_days_back=tavily_days_back,  # Look back 60 days initially
                            max_search_depth=2,    # Increase search depth
                            planner_model=planner_model,
                            planner_provider=planner_provider,
                            writer_provider="azure_openai",
                            writer_model="gpt-4o"
                ):
                    # Debug the event
                    debug_event(event, prefix="  ")
                    
                    # Check if this is the final report event
                    if isinstance(event, dict) and "compile_final_report" in event:
                        # Extract the final report content
                        if "final_report" in event.get("compile_final_report", {}):
                            report_content = event["compile_final_report"]["final_report"]
                            logger.info(f"Final report content extracted, length: {len(report_content)}")
                        elif "report" in event:
                            # Alternative location for the report
                            report_content = event["report"]
                            logger.info(f"Report content found in event['report'], length: {len(report_content)}")
                        else:
                            # Try to find any content that might be the report
                            for key, value in event.items():
                                if isinstance(value, str) and len(value) > 100:
                                    logger.info(f"Potential report content found in key: {key}, length: {len(value)}")
                                    if not report_content:  # Only use if we haven't found anything else
                                        report_content = value
                
                # If we got here without an exception, break out of the retry loop
                break
                    
            except Exception as api_error:
                retry_count += 1
                if "502 Bad Gateway" in str(api_error) and retry_count < max_retries:
                    wait_time = 10 * retry_count  # Exponential backoff: 10s, 20s, 30s
                    logger.warning(f"Tavily API error: {str(api_error)}. Retrying in {wait_time} seconds... (Attempt {retry_count}/{max_retries})")
                    await asyncio.sleep(wait_time)
                else:
                    # If it's not a 502 error or we've exhausted retries, re-raise the exception
                    raise
        
        # If no content was found, try again with a longer lookback period
        if not report_content:
            logger.warning(f"No content found with {tavily_days_back} days lookback. Trying with 365 days...")
            tavily_days_back = 365
            
            # Reset retry counter for the second attempt
            retry_count = 0
            while retry_count < max_retries:
                try:
                    # Run the research again with longer lookback
                    async for event in research_topic(
                        topic=topic,  # Use the provided company-specific topic
                        urls=[url],
                        tavily_days_back=tavily_days_back,  # Look back 365 days
                        max_search_depth=2,    # Increase search depth
                        planner_model=planner_model,
                        planner_provider=planner_provider,
                        writer_provider="azure_openai",
                        writer_model="gpt-4o"
                    ):
                        # Debug the event
                        debug_event(event, prefix="  ")
                        
                        # Check if this is the final report event
                        if isinstance(event, dict) and "compile_final_report" in event:
                            # Extract the final report content
                            if "final_report" in event.get("compile_final_report", {}):
                                report_content = event["compile_final_report"]["final_report"]
                                logger.info(f"Final report content extracted with {tavily_days_back} days lookback, length: {len(report_content)}")
                            elif "report" in event:
                                # Alternative location for the report
                                report_content = event["report"]
                                logger.info(f"Report content found in event['report'] with {tavily_days_back} days lookback, length: {len(report_content)}")
                    
                    # If we got here without an exception, break out of the retry loop
                    break
                    
                except Exception as api_error:
                    retry_count += 1
                    if "502 Bad Gateway" in str(api_error) and retry_count < max_retries:
                        wait_time = 10 * retry_count  # Exponential backoff: 10s, 20s, 30s
                        logger.warning(f"Tavily API error with extended lookback: {str(api_error)}. Retrying in {wait_time} seconds... (Attempt {retry_count}/{max_retries})")
                        await asyncio.sleep(wait_time)
                    else:
                        # If it's not a 502 error or we've exhausted retries, re-raise the exception
                        raise
                
    except Exception as e:
        logger.error(f"❌ Error researching {customer_name or url}: {str(e)}")
        
        # Create an error report
        error_report = f"# Research Error Report for {customer_name or url}\n\n"
        error_report += f"## URL: {url}\n\n"
        error_report += f"# Introduction\n\n"
        error_report += f"This report was intended to focus specifically on {company_name}, analyzing recent news about the company's operations, "
        error_report += f"strategic initiatives, and IT priorities with emphasis on their specific IT challenges, pain points, and desired outcomes.\n\n"
        error_report += f"## Error During Research\n\n"
        error_report += f"An error occurred during the research process while attempting to gather information specific to {company_name}:\n\n```\n{str(e)}\n```\n\n"
        
        # Add specific guidance for common errors
        if "502 Bad Gateway" in str(e):
            error_report += "This appears to be a temporary issue with the Tavily search API. "
            error_report += "The service may be experiencing high load or temporary outages. "
            error_report += "Please try again later when the service has recovered.\n\n"
        elif "429" in str(e) or "Too Many Requests" in str(e):
            error_report += "This appears to be a rate limiting issue. "
            error_report += "The API has received too many requests in a short period. "
            error_report += "Please try again later or reduce the frequency of requests.\n\n"
        
        error_report += "Please try again later or contact support if the issue persists."
        
        # Add generic discovery questions section for error reports
        error_report += f"\n\n## Discovery Questions for Cisco Sellers\n\n"
        error_report += f"While we couldn't retrieve specific information about {company_name} due to technical issues, here are some general discovery questions that could help start conversations:\n\n"
        error_report += f"1. **Current Infrastructure Assessment**: \"What are the biggest challenges you're currently facing with your IT infrastructure at {company_name}?\"\n\n"
        error_report += f"2. **Digital Transformation**: \"Is {company_name} currently undertaking or planning any digital transformation initiatives? What are your key priorities?\"\n\n"
        error_report += f"3. **Security Concerns**: \"How is {company_name} addressing cybersecurity challenges in today's increasingly complex threat landscape?\"\n\n"
        error_report += f"4. **Network Reliability**: \"How satisfied are you with the performance and reliability of your current network infrastructure?\"\n\n"
        error_report += f"5. **Future Planning**: \"What technology investments is {company_name} considering over the next 12-24 months to support your business objectives?\"\n\n"
        error_report += f"These questions can help start meaningful conversations and uncover specific needs that might align with Cisco's solutions."
        
        try:
            file_paths = save_markdown_report(
                url=url,
                content=error_report,
                topic=f"Error Report for {company_name}",
                customer_name=customer_name,
                customer_metadata=customer_metadata
            )
            
            if file_paths:
                if 'markdown_gcs_url' in file_paths:
                    logger.error(f"   Error report saved to: {file_paths['markdown_gcs_url']}")
                
                if 'docx_gcs_url' in file_paths:
                    logger.error(f"   Word document saved to: {file_paths['docx_gcs_url']}")
                
                if 'json_gcs_url' in file_paths:
                    logger.error(f"   JSON report saved to: {file_paths['json_gcs_url']}")
                
                return file_paths
            else:
                logger.error(f"Failed to save and upload error report for {customer_name or url}")
                return None
        except Exception as save_error:
            logger.error(f"   Could not save error report: {str(save_error)}")
            return None
    
    # If we have report content, save it
    if report_content and len(report_content) > 200:
        file_paths = save_markdown_report(
            url=url,
            content=report_content,
            topic=topic or f"Research on {company_name}",
            customer_name=customer_name,
            customer_metadata=customer_metadata
        )
        
        if file_paths:
            logger.info(f"✅ Research completed for {customer_name or url}")
            
            if 'markdown_gcs_url' in file_paths:
                logger.info(f"   Markdown report saved to: {file_paths['markdown_gcs_url']}")
            
            if 'docx_gcs_url' in file_paths:
                logger.info(f"   Word document saved to: {file_paths['docx_gcs_url']}")
            
            if 'json_gcs_url' in file_paths:
                logger.info(f"   JSON report saved to: {file_paths['json_gcs_url']}")
            
            return file_paths
        else:
            logger.error(f"Failed to save and upload report for {customer_name or url}")
            return None
    else:
        logger.warning(f"⚠️ No report content was generated for {customer_name or url}")
        logger.warning(f"   Total lookback period: {tavily_days_back} days")
        
        # Create a simple placeholder report
        simple_report = f"# Research Report for {customer_name or url}\n\n"
        simple_report += f"## URL: {url}\n\n"
        simple_report += f"# Introduction\n\n"
        simple_report += f"This report was intended to focus specifically on {company_name}, analyzing recent news about the company's operations, "
        simple_report += f"strategic initiatives, and IT priorities. However, no significant recent news was found within the specified lookback period of {tavily_days_back} days.\n\n"
        simple_report += f"# Discovery Questions for {company_name}\n\n"
        simple_report += f"Despite the lack of recent news, here are some general discovery questions that might help initiate conversations with {company_name}:\n\n"
        simple_report += f"1. **Current IT Infrastructure**: \"Can you tell me about your current IT infrastructure and any challenges you're facing?\"\n\n"
        simple_report += f"2. **Digital Transformation**: \"Is {company_name} currently undertaking or planning any digital transformation initiatives? What are your key priorities?\"\n\n"
        simple_report += f"3. **Security Concerns**: \"How is {company_name} addressing cybersecurity challenges in today's increasingly complex threat landscape?\"\n\n"
        simple_report += f"4. **Network Reliability**: \"How satisfied are you with the performance and reliability of your current network infrastructure?\"\n\n"
        simple_report += f"5. **Future Planning**: \"What technology investments is {company_name} considering over the next 12-24 months to support your business objectives?\"\n\n"
        simple_report += f"These questions can help start meaningful conversations and uncover specific needs that might align with Cisco's solutions."
        
        file_paths = save_markdown_report(
            url=url,
            content=simple_report,
            topic=f"No Recent News for {company_name}",
            customer_name=customer_name,
            customer_metadata=customer_metadata
        )
        
        if file_paths:
            logger.info(f"   Created a simple placeholder report instead")
            
            if 'markdown_gcs_url' in file_paths:
                logger.info(f"   Markdown report saved to: {file_paths['markdown_gcs_url']}")
            
            if 'docx_gcs_url' in file_paths:
                logger.info(f"   Word document saved to: {file_paths['docx_gcs_url']}")
            
            if 'json_gcs_url' in file_paths:
                logger.info(f"   JSON report saved to: {file_paths['json_gcs_url']}")
            
            return file_paths
        else:
            logger.error(f"Failed to save and upload placeholder report for {customer_name or url}")
            return None

def download_from_gcs(gcs_path, local_path):
    """Download a file from Google Cloud Storage."""
    try:
        client = get_gcs_client()
        if not client:
            logger.error("Failed to initialize GCS client")
            return False
            
        # Extract the blob path from the full GCS path
        if gcs_path.startswith("gs://"):
            # Remove the gs:// prefix and bucket name
            parts = gcs_path.replace("gs://", "").split("/", 1)
            if len(parts) < 2:
                logger.error(f"Invalid GCS path format: {gcs_path}")
                return False
            blob_path = parts[1]
        else:
            # Assume the path is already a blob path
            blob_path = gcs_path
        
        bucket = client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(blob_path)
        
        logger.info(f"Downloading file from gs://{GCS_BUCKET_NAME}/{blob_path} to {local_path}")
        blob.download_to_filename(local_path)
        
        if os.path.exists(local_path):
            logger.info(f"File successfully downloaded to {local_path}")
            return True
        else:
            logger.error(f"File download did not create the expected local file: {local_path}")
            return False
    except Exception as e:
        logger.error(f"Error downloading from GCS: {str(e)}")
        return False

def load_customer_data():
    """Load customer data from Excel file and filter for valid websites and Heartland-Gulf."""
    try:
        # Define the Excel file path
        excel_filename = "Customer Parquet top 80 select hierarchy for test.xlsx"
        
        if USE_GCS_EXCEL:
            # If using GCS, download the Excel file from GCS
            logger.info("Using Excel file from Google Cloud Storage")
            
            # Create a temporary directory for the downloaded file
            temp_dir = get_temp_dir()
            local_excel_path = os.path.join(temp_dir, excel_filename)
            
            # GCS path to the Excel file
            gcs_excel_path = f"{GCS_EXCEL_FOLDER}/{excel_filename}"
            
            # Download the file
            download_success = download_from_gcs(gcs_excel_path, local_excel_path)
            
            if not download_success:
                logger.error(f"Failed to download Excel file from GCS. Falling back to local file.")
                # Fall back to the local file
                excel_path = excel_filename
            else:
                excel_path = local_excel_path
        else:
            # Use the local Excel file
            excel_path = excel_filename
        
        # Load the Excel file
        logger.info(f"Loading customer data from: {excel_path}")
        df = pd.read_excel(excel_path)
        
        logger.info(f"Loaded {len(df)} rows from Excel file")
        
        # Filter for valid websites (not blank or ".")
        df = df[df['WEBSITE'].notna()]  # Remove NaN values
        df = df[df['WEBSITE'] != "."]   # Remove "." values
        df = df[df['WEBSITE'] != ""]    # Remove empty strings
        
        # Filter for Heartland-Gulf in Sales Level 4
        df = df[df['SALES_LEVEL_4'] == 'HEARTLAND-GULF COMMERCIAL OPERATION']
        
        logger.info(f"Filtered to {len(df)} customers with valid websites and Heartland-Gulf region")
        
        # Create a list of tuples with (website, customer_name, metadata_dict, topic)
        customer_data = []
        for _, row in df.iterrows():
            website = row['WEBSITE'].strip() if isinstance(row['WEBSITE'], str) else row['WEBSITE']
            customer_name = row['SAVM_NAME_WITH_ID'] if 'SAVM_NAME_WITH_ID' in row else None
            
            # Create a dictionary with all available metadata from the Excel row
            metadata = {}
            for column in df.columns:
                # Convert any non-serializable values to strings
                if pd.notna(row[column]):
                    if isinstance(row[column], (int, float, str, bool)):
                        metadata[column] = row[column]
                    else:
                        metadata[column] = str(row[column])
            
            # Create a company-specific topic for each URL
            company_name = extract_company_name(website, customer_name)
            company_specific_topic = f"""Research and analyze news specifically about {company_name}.

The report structure should follow this exact order:

1. Company News (REQUIRES RESEARCH)
   - Any substantial news about the specific company {company_name}
   - Recent acquisitions, mergers, or partnerships involving {company_name}
   - Leadership changes or restructuring at {company_name}

2. News about their IT Priorities (REQUIRES RESEARCH)
   - {company_name}'s specific IT challenges, pain points, and desired outcomes
   - {company_name}'s technology investments or initiatives
   - {company_name}'s digital transformation efforts and IT strategy
   - Specific IT projects, implementations, or challenges faced by {company_name}

3. Discovery Questions for Cisco Sellers
   - Based on the findings in sections 1 and 2, create 3-5 thoughtful discovery questions that Cisco sellers can use to start conversations with {company_name}
   - Questions should relate to their IT challenges, pain points, or initiatives identified in the research
   - Questions should be open-ended and designed to uncover opportunities for Cisco solutions

Note: Focus ONLY on {company_name}. Do NOT include general industry trends or news about other companies unless they directly relate to {company_name}.
IMPORTANT: All sections REQUIRE research to find specific information about {company_name}.
"""
            
            customer_data.append((website, customer_name, metadata, company_specific_topic))
        
        return customer_data
    
    except Exception as e:
        logger.error(f"Error loading customer data: {str(e)}")
        return []

async def main():
    """Main entry point for the research agent"""
    customer_data = load_customer_data()
    
    if not customer_data:
        logger.warning("No valid customers found. Exiting.")
        return
    
    logger.info(f"Starting research on {len(customer_data)} customers:")
    for i, (url, customer_name, _, _) in enumerate(customer_data[:10]):
        logger.info(f"- {customer_name or 'Unknown'}: {url}")
    
    if len(customer_data) > 10:
        logger.info(f"... and {len(customer_data) - 10} more")
    
    async def process_all_urls():
        for url, customer_name, customer_metadata, topic in customer_data:
            logger.info(f"\nProcessing URL: {url} for {customer_name or 'Unknown'}")
            
            try:
                await process_url(url, topic, customer_name, customer_metadata)
            except Exception as e:
                logger.error(f"Error processing {url}: {str(e)}")
                
                # Create an error report
                try:
                    company_name = extract_company_name(url, customer_name)
                    error_report = f"# Research Error Report for {customer_name or url}\n\n"
                    error_report += f"## URL: {url}\n\n"
                    error_report += f"# Introduction\n\n"
                    error_report += f"This report was intended to focus specifically on {company_name}, analyzing recent news about the company's operations, "
                    error_report += f"strategic initiatives, and IT priorities with emphasis on their specific IT challenges, pain points, and desired outcomes.\n\n"
                    error_report += f"## Error During Research\n\n"
                    error_report += f"An error occurred during the research process while attempting to gather information specific to {company_name}:\n\n```\n{str(e)}\n```\n\n"
                    
                    # Add specific guidance for common errors
                    if "502 Bad Gateway" in str(e):
                        error_report += "This appears to be a temporary issue with the Tavily search API. "
                        error_report += "The service may be experiencing high load or temporary outages. "
                        error_report += "Please try again later when the service has recovered.\n\n"
                    elif "429" in str(e) or "Too Many Requests" in str(e):
                        error_report += "This appears to be a rate limiting issue. "
                        error_report += "The API has received too many requests in a short period. "
                        error_report += "Please try again later or reduce the frequency of requests.\n\n"
                    
                    error_report += "Please try again later or contact support if the issue persists."
                    
                    # Add generic discovery questions section for error reports
                    error_report += f"\n\n## Discovery Questions for Cisco Sellers\n\n"
                    error_report += f"While we couldn't retrieve specific information about {company_name} due to technical issues, here are some general discovery questions that could help start conversations:\n\n"
                    error_report += f"1. **Current Infrastructure Assessment**: \"What are the biggest challenges you're currently facing with your IT infrastructure at {company_name}?\"\n\n"
                    error_report += f"2. **Digital Transformation**: \"Is {company_name} currently undertaking or planning any digital transformation initiatives? What are your key priorities?\"\n\n"
                    error_report += f"3. **Security Concerns**: \"How is {company_name} addressing cybersecurity challenges in today's increasingly complex threat landscape?\"\n\n"
                    error_report += f"4. **Network Reliability**: \"How satisfied are you with the performance and reliability of your current network infrastructure?\"\n\n"
                    error_report += f"5. **Future Planning**: \"What technology investments is {company_name} considering over the next 12-24 months to support your business objectives?\"\n\n"
                    error_report += f"These questions can help start meaningful conversations and uncover specific needs that might align with Cisco's solutions."
                    
                    # Save the error report using the existing function
                    try:
                        file_paths = save_markdown_report(
                            url=url,
                            content=error_report,
                            topic=f"Error Report for {company_name}",
                            customer_name=customer_name,
                            customer_metadata=customer_metadata
                        )
                        
                        if file_paths:
                            if 'markdown_gcs_url' in file_paths:
                                logger.error(f"   Error report saved to: {file_paths['markdown_gcs_url']}")
                            
                            if 'docx_gcs_url' in file_paths:
                                logger.error(f"   Word document saved to: {file_paths['docx_gcs_url']}")
                            
                            if 'json_gcs_url' in file_paths:
                                logger.error(f"   JSON report saved to: {file_paths['json_gcs_url']}")
                            
                            return file_paths
                        else:
                            logger.error(f"Failed to save and upload error report for {customer_name or url}")
                            return None
                    except Exception as save_error:
                        logger.error(f"   Could not save error report: {str(save_error)}")
                        return None
                except Exception as report_error:
                    logger.error(f"   Could not create error report: {str(report_error)}")
    
    await process_all_urls()
    logger.info("\nAll research tasks completed!")

if __name__ == "__main__":
    logger.info("Starting research process...")
    
    try:
        asyncio.run(main())
    except Exception as e:
        logger.error(f"Error in main execution: {str(e)}")
